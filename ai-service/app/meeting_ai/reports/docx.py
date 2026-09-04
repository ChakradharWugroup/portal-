import os
from docx import Document
from docx.shared import Pt, RGBColor

class DocxReportGenerator:
    def __init__(self, output_dir="reports"):
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def generate_report(self, meeting_title: str, summary: str, action_items: list, transcript: str) -> str:
        """
        Generates a styled Microsoft Word document and returns the file path.
        """
        safe_title = "".join([c for c in meeting_title if c.isalpha() or c.isdigit()]).rstrip()
        filename = f"{self.output_dir}/{safe_title}_Report.docx"
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f"Enterprise AI Meeting Report: {meeting_title}", level=0)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(summary)
        
        # Action Items
        if action_items:
            doc.add_heading('Action Items', level=1)
            for item in action_items:
                desc = item.get("description", "")
                owner = item.get("owner", "Unassigned")
                p = doc.add_paragraph(style='List Bullet')
                runner = p.add_run(f"{owner}: ")
                runner.bold = True
                p.add_run(desc)
                
        # Full Transcript
        doc.add_heading('Full Diarized Transcript', level=1)
        doc.add_paragraph(transcript)
        
        doc.save(filename)
        print(f"Generated DOCX: {filename}")
        return filename
